import tkinter as tk
from tkinter import messagebox, simpledialog, filedialog
import json
from pathlib import Path
import pandas as pd
from datetime import datetime, timedelta
import random
import os
import sys
import re
from docx import Document

# ------------------------ CONFIGS INICIAIS ------------------------

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return Path(sys._MEIPASS) / relative_path
    return Path(__file__).parent / relative_path

caminho_usuarios = resource_path('usuarios.json')
caminho_historico = resource_path('historico_nomes.json')
caminho_modelo_docx = resource_path('CICLO.docx')

def inicializar_arquivos():
    if not caminho_usuarios.exists():
        with open(caminho_usuarios, 'w', encoding='utf-8') as f:
            json.dump({"elismar": "266586"}, f)
    if not caminho_historico.exists():
        with open(caminho_historico, 'w', encoding='utf-8') as f:
            json.dump([], f)

def carregar_usuarios():
    with open(caminho_usuarios, 'r', encoding='utf-8') as f:
        return json.load(f)

def salvar_usuarios(usuarios):
    with open(caminho_usuarios, 'w', encoding='utf-8') as f:
        json.dump(usuarios, f, indent=4)

def carregar_historico():
    if not caminho_historico.exists() or caminho_historico.stat().st_size == 0:
        return []
    with open(caminho_historico, 'r', encoding='utf-8') as f:
        return json.load(f)

def salvar_historico(historico):
    with open(caminho_historico, 'w', encoding='utf-8') as f:
        json.dump(historico, f, ensure_ascii=False, indent=4)

# ------------------------ FUNÇÕES PRINCIPAIS ------------------------

def gerar_documento():
    arquivo_xls = filedialog.askopenfilename(
        title="Selecione o arquivo ciclo.XLS",
        filetypes=[("Arquivos Excel", "*.xls")]
    )
    if not arquivo_xls:
        messagebox.showerror("Erro", "Nenhum arquivo selecionado!")
        return

    dados = pd.read_excel(arquivo_xls, engine='xlrd')
    dados_filtrado = dados[dados['psv_cid'] == "CONSULTA NO CONSULTO"]

    nome_medico_excel = dados_filtrado['psv_apel'].iloc[0].strip()
    nome_medico_final = f"DR. {nome_medico_excel}"
    crm_numero = str(dados_filtrado['fle_psv_cod'].iloc[0])

    pacientes = dados_filtrado['pac_nome'].dropna().tolist()
    pacientes = [p.split(',')[1].strip() if ',' in p else p.strip() for p in pacientes]

    hoje = datetime.now()
    historico = carregar_historico()
    historico = [item for item in historico if (hoje - datetime.strptime(item['data'], "%Y-%m-%d")) <= timedelta(days=30)]
    bloqueados = [item['nome'] for item in historico]
    pacientes_disponiveis = [p for p in pacientes if p not in bloqueados]

    quantidade_para_sortear = min(20, len(pacientes_disponiveis))
    selecionados = random.sample(pacientes_disponiveis, quantidade_para_sortear)

    if quantidade_para_sortear < 20:
        faltaram = 20 - quantidade_para_sortear
        messagebox.showwarning("Aviso", f"Faltaram {faltaram} nomes para completar 20.")

    for nome in selecionados:
        historico.append({"nome": nome, "data": hoje.strftime("%Y-%m-%d")})
    salvar_historico(historico)

    doc = Document(caminho_modelo_docx)
    data_hoje = hoje.strftime('%d/%m/%Y')

    for paragrafo in doc.paragraphs:
        if "DR. DENISE LUCAS VIANA" in paragrafo.text:
            paragrafo.clear()
            run = paragrafo.add_run(nome_medico_final)
            run.bold = True
        if "20/06/2023" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("20/06/2023", data_hoje)
        if "CRM:" in paragrafo.text:
            paragrafo.text = re.sub(r"(CRM:\s*)\d+", f"CRM: {crm_numero}", paragrafo.text)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto = celula.text.strip()
                if "20/06/2023" in texto:
                    celula.text = texto.replace("20/06/2023", data_hoje)
                if "DR. DENISE LUCAS VIANA" in texto:
                    celula.clear()
                    run = celula.paragraphs[0].add_run(nome_medico_final)
                    run.bold = True
                if "CRM:" in texto:
                    celula.text = re.sub(r"(CRM:\s*)\d+", f"CRM: {crm_numero}", texto)

    if len(doc.tables) > 1:
        tabela_pacientes = doc.tables[1]
        for idx, nome in enumerate(selecionados):
            if idx < len(tabela_pacientes._cells):
                tabela_pacientes._cells[idx].text = nome

    desktop_path = Path.home() / "Desktop"
    desktop_path.mkdir(exist_ok=True)
    novo_docx = desktop_path / "CICLO_ATUALIZADO.docx"
    doc.save(novo_docx)

    messagebox.showinfo("Sucesso", f"Documento gerado em:\n{novo_docx}")
    os.startfile(novo_docx)

def apagar_historico():
    if messagebox.askyesno("Confirmação", "Deseja realmente apagar todo o histórico?"):
        salvar_historico([])
        messagebox.showinfo("Feito", "Histórico apagado!")

def cadastrar_usuario():
    usuarios = carregar_usuarios()
    novo_login = simpledialog.askstring("Cadastro", "Novo login:")
    nova_senha = simpledialog.askstring("Cadastro", "Nova senha:", show="*")
    if novo_login and nova_senha:
        if novo_login in usuarios:
            messagebox.showerror("Erro", "Usuário já existe!")
        else:
            usuarios[novo_login] = nova_senha
            salvar_usuarios(usuarios)
            messagebox.showinfo("Sucesso", f"Usuário {novo_login} cadastrado!")

def deletar_usuario():
    usuarios = carregar_usuarios()
    lista = [u for u in usuarios.keys() if u != "elismar"]

    if not lista:
        messagebox.showinfo("Info", "Nenhum usuário disponível para excluir!")
        return

    selecao = simpledialog.askstring(
        "Excluir Usuário",
        "Usuários disponíveis:\n" + "\n".join(lista) + "\n\nDigite o login que deseja excluir:"
    )

    if selecao and selecao in usuarios and selecao != "elismar":
        if messagebox.askyesno("Confirmação", f"Confirma excluir usuário '{selecao}'?"):
            del usuarios[selecao]
            salvar_usuarios(usuarios)
            messagebox.showinfo("Sucesso", f"Usuário {selecao} excluído!")
    else:
        messagebox.showerror("Erro", "Usuário inválido ou não encontrado.")

# ------------------------ TELAS ------------------------

def abrir_painel_usuario(usuario):
    painel = tk.Tk()
    painel.title(f"Bem-vindo, {usuario}")

    gerar_btn = tk.Button(painel, text="Gerar Documento", command=gerar_documento)
    gerar_btn.pack(pady=10)

    if usuario == "elismar":
        apagar_btn = tk.Button(painel, text="Apagar Histórico", command=apagar_historico)
        apagar_btn.pack(pady=5)

        cadastrar_btn = tk.Button(painel, text="Cadastrar Novo Usuário", command=cadastrar_usuario)
        cadastrar_btn.pack(pady=5)

        deletar_btn = tk.Button(painel, text="Deletar Usuário", command=deletar_usuario)
        deletar_btn.pack(pady=5)

    painel.mainloop()

def fazer_login():
    usuarios = carregar_usuarios()
    login = login_entry.get()
    senha = senha_entry.get()

    if login in usuarios and usuarios[login] == senha:
        messagebox.showinfo("Sucesso", f"Login bem-sucedido!\nBem-vindo {login}!")
        login_window.destroy()
        abrir_painel_usuario(login)
    else:
        messagebox.showerror("Erro", "Login ou senha incorretos!")

# ------------------------ EXECUÇÃO ------------------------

inicializar_arquivos()

login_window = tk.Tk()
login_window.title("Login - Controle de Ciclo")

tk.Label(login_window, text="Login:").pack()
login_entry = tk.Entry(login_window)
login_entry.pack()

tk.Label(login_window, text="Senha:").pack()
senha_entry = tk.Entry(login_window, show="*")
senha_entry.pack()

entrar_btn = tk.Button(login_window, text="Entrar", command=fazer_login)
entrar_btn.pack(pady=10)

login_window.mainloop()
